import os
from docx import Document

FILES = [
    r"C:\BattleBot\CLAUDE.md",
    r"C:\BattleBot\STATE.md",
    r"C:\BattleBot\ANTI-PATTERNS.md",
    r"C:\Users\Admin\.claude\projects\C--BattleBot\memory\MEMORY.md",
]

TARGET = r"C:\Users\Admin\Потоковая передача файлов Google Диска\Мой диск\Gemini_Sync"


def md_to_docx(src, dst):
    doc = Document()
    with open(src, encoding="utf-8") as f:
        for line in f:
            line = line.rstrip()
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)
    doc.save(dst)


def sync():
    os.makedirs(TARGET, exist_ok=True)
    ok = 0
    for src in FILES:
        name = os.path.splitext(os.path.basename(src))[0] + ".docx"
        dst = os.path.join(TARGET, name)
        if not os.path.exists(src):
            print(f"  [SKIP] {name}")
            continue
        try:
            md_to_docx(src, dst)
            print(f"  [OK]   {name}")
            ok += 1
        except Exception as e:
            print(f"  [ERR]  {name} -> {e}")
    print(f"\n  Готово: {ok}/{len(FILES)} файлов -> {TARGET}")


if __name__ == "__main__":
    print("=== Gemini Sync ===\n")
    sync()
